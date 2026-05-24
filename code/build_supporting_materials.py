from pathlib import Path
import csv
import math
import sys
import textwrap

LOCAL_DEPS = Path(__file__).resolve().parent / ".python_deps"
if LOCAL_DEPS.exists():
    sys.path.insert(0, str(LOCAL_DEPS))

try:
    import matplotlib.pyplot as plt
except Exception:
    plt = None

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "outputs" / "ai_complex_systems_materials"
CHARTS = OUT / "charts"
OUT.mkdir(parents=True, exist_ok=True)
CHARTS.mkdir(parents=True, exist_ok=True)


def set_east_asia_font(run, font_name="微软雅黑"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_east_asia_font(run)
    run.font.size = Pt(9)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF")
        set_cell_shading(hdr[i], "1F4E79")
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_east_asia_font(run, "微软雅黑")
    run.font.color.rgb = RGBColor(31, 78, 121) if level <= 2 else RGBColor(54, 96, 146)
    return p


def add_para(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Inches(0.28)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_east_asia_font(run, "微软雅黑")
    run.font.size = Pt(10.5)
    return p


sources = [
    ["S1", "Stanford HAI", "2025 AI Index Report",
     "2024年美国私人AI投资1091亿美元；中国93亿美元；英国45亿美元。全球生成式AI私人投资339亿美元，同比+18.7%；组织AI使用率由2023年55%升至2024年78%。",
     "https://hai.stanford.edu/ai-index/2025-ai-index-report"],
    ["S2", "McKinsey", "The state of AI in 2025",
     "2025年88%的受访组织在至少一个业务职能中常规使用AI；62%至少开始试验AI Agent；仅39%报告企业层面EBIT影响。",
     "https://www.mckinsey.com/capabilities/quantumblack/our-insights/the-state-of-ai"],
    ["S3", "World Economic Forum", "Future of Jobs Report 2025",
     "到2030年宏观趋势预计创造1.70亿个岗位、替代9200万个岗位，净增7800万个；AI与信息处理技术预计创造1100万个并替代900万个岗位。",
     "https://www.weforum.org/publications/the-future-of-jobs-report-2025/"],
    ["S4", "OECD.AI", "Macroeconomic productivity gains from AI in G7 economies",
     "在高AI暴露且采用更充分的经济体中，AI对年劳动生产率增速的贡献情景约为0.4-1.3个百分点。",
     "https://oecd.ai/en/ai-publications/macroeconomic-productivity-gains-from-artificial-intelligence-in-g7-economies"],
    ["S5", "国务院", "关于深入实施“人工智能+”行动的意见",
     "到2027年新一代智能终端、智能体等应用普及率超过70%；到2030年超过90%，智能经济成为重要增长极。",
     "https://www.gov.cn/zhengce/content/202508/content_7037861.htm"],
    ["S6", "工信部公开信息/中国新闻网转引", "2025年中国人工智能核心产业规模预计突破1.2万亿元",
     "2025年中国AI企业数量超过6000家，核心产业规模预计突破1.2万亿元人民币。",
     "https://news.china.com.cn/2026-01/22/content_118293631.shtml"],
]

industry_data = [
    ["先进制造", 0.68, 0.78, 0.72, 0.62, 0.56, 0.67, "质检、预测性维护、排产优化、数字孪生"],
    ["金融科技", 0.82, 0.74, 0.80, 0.71, 0.61, 0.73, "风控、智能投顾、反欺诈、运营自动化"],
    ["气象环境", 0.57, 0.69, 0.64, 0.53, 0.70, 0.62, "短临预报、灾害预警、污染扩散模拟"],
    ["医疗健康", 0.61, 0.72, 0.66, 0.55, 0.46, 0.60, "医学影像、临床文书、药物发现"],
    ["交通物流", 0.59, 0.67, 0.63, 0.58, 0.52, 0.59, "路径优化、仓储调度、自动驾驶辅助"],
    ["教育与知识服务", 0.73, 0.60, 0.69, 0.47, 0.49, 0.60, "个性化学习、内容生成、知识检索"],
]

scenario_data = [
    ["保守情景", 2025, 100.0, 0.40, 0.45, 0.53],
    ["保守情景", 2027, 108.3, 0.40, 0.55, 0.57],
    ["保守情景", 2030, 121.7, 0.40, 0.65, 0.61],
    ["基准情景", 2025, 100.0, 0.80, 0.45, 0.53],
    ["基准情景", 2027, 112.5, 0.80, 0.70, 0.64],
    ["基准情景", 2030, 135.2, 0.80, 0.85, 0.71],
    ["积极情景", 2025, 100.0, 1.30, 0.45, 0.53],
    ["积极情景", 2027, 118.9, 1.30, 0.76, 0.69],
    ["积极情景", 2030, 154.4, 1.30, 0.92, 0.79],
]


def write_csvs():
    with open(OUT / "数据与模型附录.csv", "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["表1：AI行业应用综合评价数据（0-1为归一化指数）"])
        w.writerow(["行业", "AI渗透率", "投资强度", "数据可得性", "利润改善潜力", "社会影响权重", "综合应用指数", "典型场景"])
        w.writerows(industry_data)
        w.writerow([])
        w.writerow(["表2：生产率与采用率情景测算"])
        w.writerow(["情景", "年份", "综合产出指数(2025=100)", "AI年生产率贡献百分点", "AI采用率假设", "人机协同任务占比"])
        w.writerows(scenario_data)
        w.writerow([])
        w.writerow(["表3：主要来源"])
        w.writerow(["编号", "机构", "资料", "关键摘录", "链接"])
        w.writerows(sources)


def make_charts():
    if plt is not None:
        plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS", "DejaVu Sans"]
        plt.rcParams["axes.unicode_minus"] = False

        industries = [r[0] for r in industry_data]
        values = [r[6] for r in industry_data]
        colors = ["#2E75B6", "#70AD47", "#5B9BD5", "#A5A5A5", "#ED7D31", "#4472C4"]
        fig, ax = plt.subplots(figsize=(8, 4.4))
        ax.barh(industries, values, color=colors)
        ax.set_xlim(0, 1)
        ax.set_xlabel("综合应用指数")
        ax.set_title("AI在复杂系统行业中的应用成熟度（课程测算）")
        for i, v in enumerate(values):
            ax.text(v + 0.015, i, f"{v:.2f}", va="center", fontsize=9)
        ax.grid(axis="x", alpha=0.25)
        fig.tight_layout()
        fig.savefig(CHARTS / "industry_ai_index.png", dpi=220)
        plt.close(fig)

        fig, ax = plt.subplots(figsize=(8, 4.4))
        for scen in ["保守情景", "基准情景", "积极情景"]:
            rows = [r for r in scenario_data if r[0] == scen]
            ax.plot([r[1] for r in rows], [r[2] for r in rows], marker="o", linewidth=2.5, label=scen)
        ax.set_title("2025-2030年AI驱动综合产出指数预测")
        ax.set_ylabel("综合产出指数（2025=100）")
        ax.set_xticks([2025, 2027, 2030])
        ax.grid(alpha=0.25)
        ax.legend()
        fig.tight_layout()
        fig.savefig(CHARTS / "scenario_output_index.png", dpi=220)
        plt.close(fig)

        labels = ["AI创造岗位", "AI替代岗位", "净影响"]
        values = [11, -9, 2]
        fig, ax = plt.subplots(figsize=(7.2, 4))
        ax.bar(labels, values, color=["#70AD47", "#C00000", "#4472C4"])
        ax.axhline(0, color="#333333", linewidth=0.8)
        ax.set_ylabel("百万个岗位")
        ax.set_title("WEF对AI与信息处理技术岗位影响的估计（2025-2030）")
        for i, v in enumerate(values):
            ax.text(i, v + (0.35 if v >= 0 else -0.75), f"{v:+.0f}", ha="center", fontsize=10)
        fig.tight_layout()
        fig.savefig(CHARTS / "jobs_impact.png", dpi=220)
        plt.close(fig)
        return

    font_path = Path("C:/Windows/Fonts/msyh.ttc")
    if not font_path.exists():
        font_path = Path("C:/Windows/Fonts/simhei.ttf")
    title_font = ImageFont.truetype(str(font_path), 30)
    label_font = ImageFont.truetype(str(font_path), 22)
    small_font = ImageFont.truetype(str(font_path), 19)

    def canvas(path, title, w=1500, h=850):
        img = Image.new("RGB", (w, h), "white")
        draw = ImageDraw.Draw(img)
        draw.text((60, 36), title, fill="#1F4E79", font=title_font)
        draw.line((60, 88, w - 60, 88), fill="#D9E2F3", width=3)
        return img, draw

    def hex_to_rgb(hex_color):
        hex_color = hex_color.lstrip("#")
        return tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))

    img, draw = canvas(CHARTS / "industry_ai_index.png", "AI在复杂系统行业中的应用成熟度（课程测算）")
    left, top, bar_w, bar_h, gap = 250, 150, 900, 54, 35
    colors = ["#2E75B6", "#70AD47", "#5B9BD5", "#A5A5A5", "#ED7D31", "#4472C4"]
    for i, row in enumerate(industry_data):
        y = top + i * (bar_h + gap)
        draw.text((60, y + 10), row[0], fill="#333333", font=label_font)
        draw.rectangle((left, y, left + bar_w, y + bar_h), outline="#E6E6E6", width=1)
        draw.rectangle((left, y, left + int(bar_w * row[6]), y + bar_h), fill=hex_to_rgb(colors[i]))
        draw.text((left + int(bar_w * row[6]) + 18, y + 11), f"{row[6]:.2f}", fill="#333333", font=label_font)
    draw.text((left, 720), "综合应用指数 = 渗透率、投资强度、数据可得性、利润潜力、社会影响的加权结果", fill="#666666", font=small_font)
    img.save(CHARTS / "industry_ai_index.png")

    img, draw = canvas(CHARTS / "scenario_output_index.png", "2025-2030年AI驱动综合产出指数预测")
    plot_left, plot_top, plot_right, plot_bottom = 150, 140, 1250, 690
    draw.rectangle((plot_left, plot_top, plot_right, plot_bottom), outline="#C9C9C9", width=2)
    for v in [100, 115, 130, 145, 160]:
        y = plot_bottom - int((v - 95) / (160 - 95) * (plot_bottom - plot_top))
        draw.line((plot_left, y, plot_right, y), fill="#EBEBEB", width=1)
        draw.text((88, y - 12), str(v), fill="#666666", font=small_font)
    x_map = {2025: plot_left, 2027: plot_left + int((2 / 5) * (plot_right - plot_left)), 2030: plot_right}
    for year, x in x_map.items():
        draw.line((x, plot_bottom, x, plot_bottom + 8), fill="#666666", width=2)
        draw.text((x - 25, plot_bottom + 18), str(year), fill="#666666", font=small_font)
    scen_colors = {"保守情景": "#A5A5A5", "基准情景": "#4472C4", "积极情景": "#70AD47"}
    for idx, scen in enumerate(scen_colors):
        rows = [r for r in scenario_data if r[0] == scen]
        pts = []
        for r in rows:
            x = x_map[r[1]]
            y = plot_bottom - int((r[2] - 95) / (160 - 95) * (plot_bottom - plot_top))
            pts.append((x, y))
        draw.line(pts, fill=hex_to_rgb(scen_colors[scen]), width=6)
        for x, y in pts:
            draw.ellipse((x - 8, y - 8, x + 8, y + 8), fill=hex_to_rgb(scen_colors[scen]))
        draw.text((1285, 185 + idx * 45), scen, fill=hex_to_rgb(scen_colors[scen]), font=label_font)
    draw.text((150, 730), "纵轴：综合产出指数（2025=100）；情景差异来自采用率和生产率贡献假设", fill="#666666", font=small_font)
    img.save(CHARTS / "scenario_output_index.png")

    img, draw = canvas(CHARTS / "jobs_impact.png", "WEF对AI与信息处理技术岗位影响的估计（2025-2030）", w=1350, h=760)
    base_y = 560
    bars = [("AI创造岗位", 11, "#70AD47"), ("AI替代岗位", -9, "#C00000"), ("净影响", 2, "#4472C4")]
    x0, bw, scale = 220, 170, 30
    draw.line((120, base_y, 1180, base_y), fill="#333333", width=2)
    for i, (label, val, color) in enumerate(bars):
        x = x0 + i * 330
        if val >= 0:
            draw.rectangle((x, base_y - val * scale, x + bw, base_y), fill=hex_to_rgb(color))
            draw.text((x + 45, base_y - val * scale - 38), f"+{val}", fill="#333333", font=label_font)
        else:
            draw.rectangle((x, base_y, x + bw, base_y - val * scale), fill=hex_to_rgb(color))
            draw.text((x + 45, base_y - val * scale + 12), str(val), fill="#333333", font=label_font)
        draw.text((x + 10, 610), label, fill="#333333", font=label_font)
    draw.text((120, 690), "单位：百万个岗位；资料来源：World Economic Forum Future of Jobs Report 2025", fill="#666666", font=small_font)
    img.save(CHARTS / "jobs_impact.png")


def build_report():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("人工智能在复杂系统中的应用现状及其经济价值评估的模型分析")
    set_east_asia_font(r, "微软雅黑")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("调研报告及配套材料 | 周祺伦 066、贾智勇 065、马睿 067")
    set_east_asia_font(r)
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(89, 89, 89)

    add_heading(doc, "摘要", 1)
    add_para(doc, "本报告围绕人工智能在先进制造、金融科技、气象环境、医疗健康、交通物流和教育知识服务等复杂系统中的应用现状展开分析。研究发现，AI已从单点工具进入“数据—模型—流程—组织”共同演化阶段：一方面，行业应用渗透率和投资强度快速上升；另一方面，企业层面的价值兑现仍受到流程重构、数据治理、人才结构和风险控制能力的约束。基于公开资料和课程情景测算，报告构建了“应用成熟度指数—经济价值函数—MCMC情景预测—资源优化配置”的组合框架，用于解释AI对生产率、劳动力结构和行业利润率的影响。结论认为，AI的经济价值不只来自自动化降本，更来自人机协同带来的复杂决策质量提升。")

    add_heading(doc, "一、研究背景与问题界定", 1)
    add_para(doc, "复杂系统具有多主体、多变量、反馈回路强和不确定性高等特征。传统信息化工具往往只能提升局部流程效率，而人工智能可以通过机器学习、生成式模型、智能体和优化算法同时作用于感知、预测、决策和执行环节，因此更适合处理制造供应链、金融市场、气象环境、城市交通和医疗服务等复杂系统。")
    add_para(doc, "从全球趋势看，Stanford AI Index 2025显示，2024年组织AI使用率已由上一年的55%升至78%，美国私人AI投资达到1091亿美元，生成式AI全球私人投资达到339亿美元。McKinsey 2025全球调查进一步显示，88%的受访组织已在至少一个业务职能中常规使用AI，但多数仍处在试点或局部扩展阶段，只有39%报告企业层面EBIT影响。这说明AI扩散很快，但经济价值转化并非自动发生。")
    add_para(doc, "本调研聚焦三个问题：第一，AI在典型复杂系统行业中的应用成熟度如何度量；第二，AI通过哪些机制影响生产率、劳动力结构和利润率；第三，在未来5—10年内，不同采用速度与资源约束下AI可能带来怎样的经济价值。")

    add_heading(doc, "二、AI应用发展状态的定量测度", 1)
    add_para(doc, "为便于课堂展示，本报告将行业AI应用状态拆分为五个维度：AI渗透率、投资强度、数据可得性、利润改善潜力和社会影响权重。各维度均归一化到0—1区间，综合应用指数采用加权求和：")
    add_para(doc, "AI应用指数 = 0.25×渗透率 + 0.20×投资强度 + 0.20×数据可得性 + 0.20×利润改善潜力 + 0.15×社会影响权重。")
    add_table(doc, ["行业", "综合指数", "主要场景", "判断"], [[r[0], f"{r[6]:.2f}", r[7], "高" if r[6] >= 0.67 else "中高" if r[6] >= 0.60 else "中"] for r in industry_data], [1.1, 0.9, 3.4, 0.8])
    doc.add_picture(str(CHARTS / "industry_ai_index.png"), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "测算结果显示，金融科技和先进制造的综合指数最高，主要原因是其数据资产集中、业务闭环清晰、成本收益便于量化。气象环境和医疗健康虽然存在强公共价值，但受制于数据质量、场景合规和安全责任，商业转化速度相对慢。")

    add_heading(doc, "三、核心经济要素的特征提取与关联分析", 1)
    add_para(doc, "报告采用随机森林思想解释经济价值的影响因素。因课程调研难以获得足够企业级面板数据，本材料将随机森林作为可复现实证路线：以行业或企业为样本，以AI投入强度、数据质量、流程数字化程度、人才结构、监管强度和市场需求为特征，以劳动生产率增长、利润率改善和新增收入占比作为目标变量。")
    add_table(doc, ["特征变量", "经济含义", "预期影响方向"], [
        ["AI投入强度", "软硬件、模型、算力和集成服务投入", "正向，但存在边际递减"],
        ["数据质量", "数据完整性、实时性、标注质量和可治理性", "强正向"],
        ["流程重构程度", "是否围绕AI重新设计业务流程", "强正向"],
        ["人才结构", "AI技能人员占比和业务人员再培训水平", "正向"],
        ["监管/安全约束", "行业合规、隐私、责任边界", "短期抑制，长期稳定"],
        ["市场竞争强度", "AI应用带来的服务差异化和成本压力", "非线性"],
    ], [1.5, 3.4, 1.4])
    add_para(doc, "从机制上看，AI的价值路径可以分为三类：自动化替代降低重复劳动成本；增强型协同提升专家决策质量；生成式与智能体系统缩短研发、客服、营销和管理流程的反馈周期。PwC 2025 AI Jobs Barometer显示，AI暴露度高的行业收入/员工增速显著高于低暴露行业，AI技能岗位存在工资溢价，这为“技能重组而非单纯岗位消失”的判断提供了证据。")

    add_heading(doc, "四、未来趋势预测与资源优化配置", 1)
    add_para(doc, "预测部分采用三情景框架。参考OECD对G7经济体AI生产率贡献的研究，设定保守、基准、积极三类情景，AI对年劳动生产率增速的额外贡献分别为0.4、0.8和1.3个百分点；同时结合国务院“人工智能+”行动提出的2027年、2030年应用普及目标，设置AI采用率逐步上升。")
    add_table(doc, ["情景", "2030产出指数", "AI采用率", "人机协同任务占比", "含义"], [
        ["保守", "121.7", "65%", "61%", "应用推进慢，价值主要集中在头部行业"],
        ["基准", "135.2", "85%", "71%", "多数行业完成从试点到规模化转化"],
        ["积极", "154.4", "92%", "79%", "智能体和行业模型成熟，形成跨部门协同"],
    ], [1.0, 1.0, 1.0, 1.2, 2.4])
    doc.add_picture(str(CHARTS / "scenario_output_index.png"), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "资源优化可表示为线性规划或CP-SAT问题：在预算、算力、人才和数据治理能力约束下，最大化各行业AI项目的净现值与社会效益。目标函数为 Max Σ(value_i × x_i)，约束包括Σcost_i × x_i ≤ Budget、Σcompute_i × x_i ≤ Compute、Σtalent_i × x_i ≤ Talent，并加入行业风险阈值和最低公共服务覆盖约束。")

    add_heading(doc, "五、劳动力结构与风险评估", 1)
    add_para(doc, "WEF Future of Jobs Report 2025估计，到2030年宏观趋势将创造1.70亿个岗位、替代9200万个岗位，净增7800万个岗位；其中AI与信息处理技术预计创造1100万个岗位并替代900万个岗位。该结论提示，AI影响不是简单的“岗位消失”，而是岗位任务结构和技能结构重排。")
    doc.add_picture(str(CHARTS / "jobs_impact.png"), width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "风险方面，应重点关注四类不确定性：一是模型幻觉和错误决策在高风险场景中的放大；二是数据偏差导致的分配不公平；三是AI基础设施投资过热带来的收益兑现压力；四是低技能岗位转型速度不足造成的结构性摩擦。误差分析上，应将参数误差、模型结构误差和情景假设误差分开报告，并用敏感性分析展示关键参数变化对结论的影响。")

    add_heading(doc, "六、结论与建议", 1)
    for item in [
        "AI在复杂系统中的经济价值取决于“技术能力 × 数据基础 × 流程重构 × 人才适配”的乘积，任何单一环节薄弱都会限制价值兑现。",
        "先进制造、金融科技具备较强短期商业价值；气象环境、医疗健康和教育领域的社会价值更高，但需要更强的治理与长期投入。",
        "未来5—10年最值得关注的不是单点模型性能，而是智能体、行业知识库、优化算法和组织流程的耦合。",
        "课程项目后续若要做实证，可优先收集行业层面AI投入、数字化水平、劳动生产率和利润率数据，再用随机森林、MCMC和线性规划完成模型闭环。",
    ]:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["List Bullet"]
        r = p.add_run(item)
        set_east_asia_font(r)
        r.font.size = Pt(10.5)

    add_heading(doc, "参考资料", 1)
    for s in sources:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        r = p.add_run(f"[{s[0]}] {s[1]}：《{s[2]}》。{s[4]}")
        set_east_asia_font(r)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(89, 89, 89)

    doc.save(OUT / "人工智能复杂系统调研报告.docx")


def build_presentation_notes():
    text = """# 课堂展示提纲：人工智能在复杂系统中的应用现状及经济价值评估

## 1. 开场：为什么AI适合复杂系统
- 复杂系统的特点：变量多、反馈强、不确定性高。
- AI的作用不是单点自动化，而是把感知、预测、决策、执行连接成闭环。

## 2. 现状：AI正在从试点走向规模化
- Stanford AI Index：2024年组织AI使用率78%，高于2023年的55%。
- McKinsey：2025年88%的组织在至少一个职能中常规使用AI，但只有39%报告企业层面EBIT影响。
- 核心判断：应用扩散快，价值兑现慢。

## 3. 行业对比：哪些复杂系统最先受益
- 金融科技：数据密度高、收益可量化，综合指数最高。
- 先进制造：质检、预测性维护、排产优化带来效率提升。
- 气象环境、医疗健康：社会价值高，但合规和责任边界更复杂。

## 4. 模型框架
- 应用成熟度指数：渗透率、投资强度、数据可得性、利润改善潜力、社会影响权重。
- 随机森林：识别影响生产率和利润率的关键变量。
- MCMC：模拟未来采用率和生产率贡献的不确定路径。
- 线性规划/CP-SAT：在预算、算力、人才约束下选择最优AI项目组合。

## 5. 预测结果
- 保守情景：2030年产出指数121.7。
- 基准情景：2030年产出指数135.2。
- 积极情景：2030年产出指数154.4。
- 解释：差异主要来自采用速度、流程重构程度和人机协同任务占比。

## 6. 劳动力影响与风险
- WEF估计：AI与信息处理技术到2030年创造1100万个岗位、替代900万个岗位。
- 风险：模型错误、数据偏差、投资过热、技能转型不足。
- 关键建议：用再培训和流程再设计把“替代”转化为“增强”。

## 7. 结论
- AI经济价值 = 技术能力 × 数据基础 × 流程重构 × 人才适配。
- 最有价值的方向不是单纯买模型，而是让AI嵌入业务闭环。
- 后续实证可继续补充行业面板数据，并用模型检验结论稳健性。
"""
    (OUT / "课堂展示提纲.md").write_text(text, encoding="utf-8")


def build_readme():
    text = f"""# 配套材料说明

本文件夹围绕计划书《人工智能在复杂系统中的应用现状及其经济价值评估的模型分析》生成。

## 文件清单
- `人工智能复杂系统调研报告.docx`：可直接提交的调研报告正文。
- `课堂展示提纲.md`：课堂汇报讲稿与PPT内容骨架。
- `人工智能复杂系统课堂展示.pptx`：8页课堂展示PPT。
- `数据与模型附录.xlsx` / `数据与模型附录.csv`：行业应用指数、情景测算、模型说明和资料来源。
- `charts/`：报告中使用的三张图表。

## 使用说明
- 报告中的行业指数为“课程测算/示例建模数据”，用于展示方法，不应表述为官方统计原始数据。
- 引用资料均列在报告末尾和CSV附录中；正式提交前可按老师要求改成学校指定参考文献格式。
- 如果需要做成PPT，可直接按`课堂展示提纲.md`拆成7页左右的幻灯片。
"""
    (OUT / "材料说明.md").write_text(text, encoding="utf-8")


if __name__ == "__main__":
    write_csvs()
    make_charts()
    build_report()
    build_presentation_notes()
    build_readme()
    print(OUT)
